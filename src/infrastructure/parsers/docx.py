"""DOCX text extraction using python-docx."""

import logging
from pathlib import Path

from infrastructure.parsers.exceptions import ParseError

logger = logging.getLogger(__name__)


def parse_docx(file_path: str) -> str:
    """Extract text from a DOCX file.

    :param file_path: Local path to the DOCX file.
    :returns: Extracted text (all paragraphs concatenated).
    :raises ParseError: If file cannot be read or parsed.
    """
    path = Path(file_path)
    if not path.exists():
        raise ParseError(f"DOCX file not found: {file_path}")

    try:
        from docx import Document  # lazy import — optional dependency

        doc = Document(str(path))
        text_parts: list[str] = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text.strip())

        # Also extract text from tables (preserve table structure)
        for table in doc.tables:
            table_lines: list[str] = []
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    table_lines.append(row_text)
            if table_lines:
                text_parts.append("\n".join(table_lines))

        if not text_parts:
            raise ParseError(f"DOCX contains no extractable text: {file_path}")

        return "\n\n".join(text_parts)

    except ImportError as e:
        raise ParseError(
            "python-docx is not installed. Add it to pyproject.toml."
        ) from e
    except Exception as e:
        raise ParseError(f"Failed to parse DOCX {file_path}: {e}") from e
